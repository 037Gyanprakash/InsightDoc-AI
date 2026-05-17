import os
import re
import pandas as pd
import pdfplumber
import docx
import pytesseract
import sqlite3
from PIL import Image
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document


class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )

    
    # ENTITY EXTRACTION
    
    def extract_entity(self, text: str, filename: str) -> str:
        """
        Robust entity extraction:
        1. Try from document text
        2. Fallback to filename (always works)
        """

        if text:
            lines = [l.strip() for l in text.splitlines() if l.strip()]

            if lines:
                first_line = lines[0]
                if 2 <= len(first_line.split()) <= 4:
                    return self._normalize(first_line)

            for line in lines[:15]:
                match = re.search(r"(name\s*[:\-]\s*)([A-Za-z ]+)", line, re.I)
                if match:
                    return self._normalize(match.group(2))

            name_pattern = re.compile(r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})\b")
            for line in lines[:20]:
                match = name_pattern.search(line)
                if match:
                    return self._normalize(match.group(1))

                
        base = os.path.splitext(filename)[0]
        return self._normalize(base)

    def _normalize(self, text: str) -> str:
        """
        Normalize entity for vector DB filtering.
        CRITICAL FIX: Replaces non-alpha chars with SPACE, not empty string.
        """
        text = re.sub(r"[^a-zA-Z ]", " ", text)
        text = re.sub(r"\s+", " ", text)
        return text.strip().lower()

    
    # MAIN PROCESSOR
    
    def process_file(self, file_path: str, file_type: str, original_filename: str):
        text_content = ""

        try:
            
            # PDF
            
            if file_type == "application/pdf" or file_path.endswith(".pdf"):
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text_content += (page.extract_text() or "") + "\n"

            
            # DOCX
            
            elif (
                file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                or file_path.endswith(".docx")
            ):
                doc = docx.Document(file_path)
                text_content = "\n".join(p.text for p in doc.paragraphs)

            
            # TXT
        
            elif file_type == "text/plain" or file_path.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text_content = f.read()

            
            # IMAGE (OCR)
            
            elif file_type.startswith("image/") or file_path.endswith((".jpg", ".png", ".jpeg")):
                image = Image.open(file_path)
                text_content = pytesseract.image_to_string(image)

            
            # CSV
            
            elif file_type == "text/csv" or file_path.endswith(".csv"):
                df = pd.read_csv(file_path)
                text_content = df.to_string(index=False)

            
            # SQLITE
            
            elif file_path.endswith((".db", ".sqlite")):
                conn = sqlite3.connect(file_path)
                cursor = conn.cursor()
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
                tables = cursor.fetchall()

                for table in tables:
                    table_name = table[0]
                    text_content += f"\nTable: {table_name}\n"
                    df = pd.read_sql_query(f"SELECT * FROM {table_name}", conn)
                    text_content += df.to_string(index=False) + "\n"

                conn.close()

            else:
                return []

            if not text_content.strip():
                return []

            
            # ENTITY DETECTION
            
            entity = self.extract_entity(text_content, original_filename)
            print(f"[Processor] Extracted Entity: {entity}")

            base_doc = Document(
                page_content=text_content,
                metadata={
                    "source": original_filename,
                    "entity": entity 
                }
            )

            
            # CHUNKING
            
            return self.text_splitter.split_documents([base_doc])

        except Exception as e:
            print(f"[DocumentProcessor Error] {e}")
            return []
